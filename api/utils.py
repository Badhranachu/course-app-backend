# api/utils.py

from docx import Document
from docx.shared import Pt
import os
from django.conf import settings
from django.utils import timezone
from datetime import timedelta
from docx2pdf  import convert 

import pythoncom

from api.models import Enrollment, CertificateSequence, StudentProfile
def convert_docx_to_pdf(docx_path, pdf_path):
    """
    Safely convert DOCX → PDF using Word (Windows only)
    """
    pythoncom.CoInitialize()
    try:
        convert(docx_path, pdf_path)
    finally:
        pythoncom.CoUninitialize()

def generate_certificate(*, user, course):

    # ===============================
    # STUDENT DETAILS (SAFE)
    # ===============================
    try:
        profile = user.student_profile
        name = profile.full_name.strip()
        title = "Mr. " if profile.gender == "male" else "Ms. "
    except StudentProfile.DoesNotExist:
        name = user.email.split("@")[0]
        title = "Mr. "

    # ===============================
    # ENROLLMENT DATES
    # ===============================
    enrollment = Enrollment.objects.get(user=user, course=course)
    start_date = enrollment.enrolled_at.date()
    end_date = start_date + timedelta(days=30)

    # ===============================
    # DATE + REF NUMBER
    # ===============================
    today = timezone.now().strftime("%d/%m/%Y")

    seq, _ = CertificateSequence.objects.get_or_create(id=1)
    seq.last_number += 1
    seq.save()

    ref_no = f"NEX/INT/2025/{str(seq.last_number).zfill(2)}"

    # ===============================
    # LOAD TEMPLATE
    # ===============================
    template_path = os.path.join(
        settings.BASE_DIR,
        "api",
        "static",
        "certificates",
        "template.docx"
    )

    doc = Document(template_path)

    def add_paragraph(text, bold=False, size=12):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return p

    # ===============================
    # CERTIFICATE CONTENT
    # ===============================
    add_paragraph(f"Date: {today}", size=11)
    add_paragraph(f"Ref: {ref_no}", size=11)

    add_paragraph("")
    add_paragraph("TO WHOMSOEVER IT MAY CONCERN", bold=True, size=14)
    add_paragraph("")

    add_paragraph(
        f"This is to certify that {title}{name} has successfully completed "
        f"his internship program in {course.title} from "
        f"{start_date.strftime('%d %B %Y')} to {end_date.strftime('%d %B %Y')} "
        f"at Walnex.",
        size=12
    )

    add_paragraph("")
    add_paragraph(
        "During the internship period, we found him to be extremely inquisitive, "
        "hardworking, and disciplined. He demonstrated strong interest and curiosity "
        "in understanding the functions of our core development process and consistently "
        "put in dedicated effort. He showed willingness to dive deep into both backend "
        "and frontend concepts to strengthen his practical understanding.",
        size=12
    )

    add_paragraph("")
    add_paragraph(
        "We appreciate his enthusiasm and commitment throughout the internship "
        "and wish him every success in his future endeavors.",
        size=12
    )

    # ===============================
    # SAVE DOCX
    # ===============================
    output_dir = os.path.join(settings.MEDIA_ROOT, "certificates")
    os.makedirs(output_dir, exist_ok=True)

    docx_filename = f"certificate_{user.id}_{course.id}.docx"
    pdf_filename = f"certificate_{user.id}_{course.id}.pdf"

    docx_path = os.path.join(output_dir, docx_filename)
    pdf_path = os.path.join(output_dir, pdf_filename)

    doc.save(docx_path)

    # ===============================
    # CONVERT → PDF (SAFE)
    # ===============================
    convert_docx_to_pdf(docx_path, pdf_path)

    # OPTIONAL: remove DOCX
    os.remove(docx_path)

    return pdf_path, ref_no


from django.core.mail import send_mail

def send_otp_email(email: str, otp: str):
    """
    Sends OTP email to user.
    """

    subject = "Your OTP Verification Code"
    message = (
        f"Dear User,\n\n"
        f"Your One-Time Password (OTP) is: {otp}\n\n"
        f"This OTP is valid for 5 minutes.\n"
        f"Please do not share this OTP with anyone.\n\n"
        f"Regards,\n"
        f"Nexston / Walnex Team"
    )

    send_mail(
        subject=subject,
        message=message,
        from_email=settings.EMAIL_HOST_USER,
        recipient_list=[email],
        fail_silently=False,
    )





# api/utils.py

import os
from datetime import timedelta
from django.utils import timezone
from django.core.files.base import ContentFile
from django.core.mail import EmailMessage
from django.core.files.storage import default_storage

from api.models import (
    PreCertificate,
    Certificate,
    Enrollment,
    StudentProfile
)
import logging

logger = logging.getLogger(__name__)

logger = logging.getLogger(__name__)


def delayed_transfer_and_email(precert_id):
    """
    Safe to retry multiple times.
    Sends mail ONLY after 30 days from PAYMENT DATE.
    """

    try:
        precert = PreCertificate.objects.select_related(
            "user", "course"
        ).get(id=precert_id)
    except PreCertificate.DoesNotExist:
        return

    user = precert.user
    course = precert.course

    # ---------- CHECK ENROLLMENT ----------
    try:
        enrollment = Enrollment.objects.get(
            user=user,
            course=course,
            status="completed"  # ✅ IMPORTANT
        )
    except Enrollment.DoesNotExist:
        logger.info(f"No completed enrollment: {user.email}")
        return

    # ---------- PAYMENT DATE CHECK ----------
    if not enrollment.payment_date:
        logger.warning(f"Payment date missing: {user.email}")
        return

    # 🔁 CHANGE TO days=30 IN PRODUCTION
    eligible_at = enrollment.payment_date + timedelta(minutes=2)

    if timezone.now() < eligible_at:
        logger.info(f"Not eligible yet: {user.email}")
        return

    old_path = precert.certificate_file.name
    filename = os.path.basename(old_path)

    # ---------- SEND EMAIL ----------
    try:
        name = (
            user.student_profile.full_name
            if hasattr(user, "student_profile")
            else user.email.split("@")[0]
        )

        email = EmailMessage(
            subject=f"Certificate for {course.title}",
            body=(
                f"Hi {name},\n\n"
                f"Your internship certificate is attached.\n\n"
                f"Regards,\nWalnex / Nexston"
            ),
            to=[user.email],
        )

        with default_storage.open(old_path, "rb") as f:
            email.attach(filename, f.read(), "application/pdf")

        email.send(fail_silently=False)
        logger.info(f"Email sent: {user.email}")

    except Exception:
        logger.exception("Email failed. Will retry.")
        return  # ❌ DO NOT MOVE FILE

    # ---------- MOVE FILE ----------
    try:
        new_path = f"certificates/{filename}"

        with default_storage.open(old_path, "rb") as f:
            default_storage.save(new_path, ContentFile(f.read()))

        default_storage.delete(old_path)

        Certificate.objects.update_or_create(
            user=user,
            course=course,
            defaults={
                "github_link": precert.github_link,
                "certificate_file": new_path,
            },
        )

        precert.delete()
        logger.info(f"Certificate finalized: {user.email}")

    except Exception:
        logger.exception("File move failed AFTER email")